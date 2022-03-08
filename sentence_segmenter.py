import codecs
import os

import chardet
import spacy
import utils
import docx
from docx.enum.text import WD_COLOR_INDEX
import collections
import string
import matplotlib.pyplot as plt
import pickle
from luima_sbd import sbd_utils
import glob
from argparse import ArgumentParser
import time

parser = ArgumentParser(description="Sentence Segmenter")
parser.add_argument("-f", dest="filepath")

eng_dict = spacy.load("en_core_web_md")
highlighted_doc_filepath_spacy = "./highlighted_doc/spacy/"
highlighted_doc_filepath_anno = "./highlighted_doc/anno/"


def generate_highlighted_doc(filepath: string, sentences: list, color: WD_COLOR_INDEX):
    highlighted_doc = docx.Document()
    for i in sentences:
        highlighted_doc.add_paragraph().add_run(i).font.highlight_color = color
        highlighted_doc.add_paragraph().add_run(
            "<EOS>"
        ).font.highlight_color = WD_COLOR_INDEX.GRAY_50
    highlighted_doc.save(filepath + ".docx")


def analyze_error(gt: dict, pred: dict):
    error_dict = {}
    for (k1, v1), (k2, v2) in zip(gt.items(), pred.items()):
        fn = 0
        fp = 0
        tp = len(v1)
        for index, value in enumerate(v1):
            if index < len(v2):
                if abs(value - v2[index]) > 3:
                    fn = fn + 1

        for index, value in enumerate(v2):
            if index < len(v1):
                if abs(value - v1[index]) > 3:
                    fp = fp + 1

        precision = tp / (tp + fp)
        recall = tp / (tp + fn)
        f1 = (2 * precision * recall) / (precision + recall)
        error_dict[k1] = [precision, recall, f1]

    return error_dict


def get_annoatations_dict():
    pass


def luima_unlabelled_sent_seg():
    filepath = "/home/skatoosh/project/legal_ds_project/legal-ds-project/data/unlabeled/unlabeled/"
    luima_sent_seg = []
    total = len(os.listdir(filepath)) + 1
    print("Luima Sentence Segmentation Initiating...")
    ts = time.time()
    for i in range(1,7):
        luima_sent_seg_part = []
        pkl_name = "/home/skatoosh/project/legal_ds_project/legal-ds-project/" + "luima_unlabelled_segpart" + str(i) + ".pkl"
        with open(pkl_name, "rb") as f:
            luima_sent_seg_part = pickle.load(f)
        luima_sent_seg = luima_sent_seg + luima_sent_seg_part

    print(time.time()-ts)
    # for filename in glob.glob(filepath + "*.txt"):
    #     total = total - 1
    #     if total % 100 == 0:
    #         print("Remaining files", total)
    #         print(round(time.time() - ts))
    #     doc_sent_dict = {}
    #     raw = open(filename, "rb").read()
    #     enc = chardet.detect(raw)["encoding"]
    #     if enc == "ISO-8859-1":
    #         doc_sent_dict["filename"] = os.path.basename(filename)
    #         doc_sent_dict["plainText"] = raw
    #         doc_sent_dict["encoding"] = enc
    #         with codecs.open(filename, mode="r", encoding=enc) as f:
    #             text = f.read()
    #             sentence, indices = sbd_utils.text2sentences(text)
    #             doc_sent_dict["sentences"] = sentence
    #             doc_sent_dict["startPos"] = [indice[0] for indice in indices]
    #             doc_sent_dict["endPos"] = [indice[1] for indice in indices]
    #             luima_sent_seg.append(doc_sent_dict)
    #     else:
    #         print(os.path.basename(os.path.normpath(filepath)),os.path.basename(filename), "\t Encoding:", enc)
    #
    # print("Sentence Segmentation Done. Writing Pickle file")
    # part_dict_name = (
    #     "luima_unlabelled_seg" + os.path.basename(os.path.normpath(filepath)) + ".pkl"
    # )
    # with open(part_dict_name, "wb") as f:
    #     pickle.dump(luima_sent_seg, f)
    # print("Written Pickle file!!", part_dict_name)


def luima_sbd_sent_seg():
    train_granted, train_denied, _, _, _, _ = utils.split_data()
    data, annotated_docs = utils.get_data_and_annotations()
    docid_sent_dict = {}
    docid_sent_start_pos = {}
    print("Creating Sentence Dict Using LUIMA SBD...")
    if os.path.isfile("docid_sent_dict_luima.pkl"):
        with open("docid_sent_dict_luima.pkl", "rb") as f:
            docid_sent_dict = pickle.load(f)
        with open("docid_sent_start_pos_luima.pkl", "rb") as f:
            docid_sent_start_pos = pickle.load(f)
    else:
        for i in train_granted:
            s1 = []
            start_positions = []
            sentence, indices = sbd_utils.text2sentences(i["plainText"])
            s1.append(sentence)
            start_positions.append([indice[0] for indice in indices])
            docid_sent_dict[i["_id"]] = list(s1)
            docid_sent_start_pos[i["_id"]] = start_positions[0]
        _, _, anno_dict, anno_start_pos = standard_sent_seg()

    error_dict = analyze_error(anno_start_pos, docid_sent_start_pos)


def standard_sent_seg():

    train_granted, train_denied, _, _, _, _ = utils.split_data()
    data, annotated_docs = utils.get_data_and_annotations()
    docid_sent_dict = {}
    docid_sent_start_pos = {}
    print("Creating Sentence Dict Using Spacy...")
    if os.path.isfile("docid_sent_dict.pkl"):
        with open("docid_sent_dict.pkl", "rb") as f:
            docid_sent_dict = pickle.load(f)
        with open("docid_sent_start_pos.pkl", "rb") as f:
            docid_sent_start_pos = pickle.load(f)
    else:

        for i in train_granted:
            doc = eng_dict(i["plainText"])
            s1 = []
            start_positions = []
            for sentence in doc.sents:
                s1.append(sentence.text)
                start_positions.append(sentence.start)
            docid_sent_dict[i["_id"]] = s1
            docid_sent_start_pos[i["_id"]] = start_positions

        with open("docid_sent_dict.pkl", "wb") as f:
            pickle.dump(docid_sent_dict, f)
        with open("docid_sent_start_pos.pkl", "wb") as f:
            pickle.dump(docid_sent_start_pos, f)

    NUM_HIGHLIGHTED_DOCS = len(docid_sent_dict)
    # Highlight Spacy obtained sentences
    print("Creating Highlighted Docs...")
    if len(os.listdir(highlighted_doc_filepath_spacy)) == 0:
        count = 0
        for k, v in docid_sent_dict.items():
            if count <= NUM_HIGHLIGHTED_DOCS:
                filepath = highlighted_doc_filepath_spacy + str(k)
                generate_highlighted_doc(filepath, v, WD_COLOR_INDEX.YELLOW)
                count = count + 1
    print("Done!")
    # Create Ground Truth Highlights
    print("Creating Annotations-Sentence Dict...")

    anno_list = []
    docid_anno_dict = {}
    docid_anno_start_pos = {}
    if os.path.isfile("docid_anno_dict.pkl"):
        with open("docid_anno_dict.pkl", "rb") as f:
            docid_anno_dict = pickle.load(f)
        with open("docid_anno_start_pos.pkl", "rb") as f:
            docid_anno_start_pos = pickle.load(f)
    else:
        for k, v in docid_sent_dict.items():
            for i in data["annotations"]:
                if k == i["document"]:
                    anno_list.append(i)
            docid_anno_dict[k] = anno_list
            anno_list = []

        # Changing the docid_anno_dict to document_id:[sentences] similar to docid_sent_dict
        anno_sent = {}
        anno_sent_dict = {}
        tgcnt = 0

        for k, v in docid_anno_dict.items():
            for i in v:
                anno_sent_dict[i["start"]] = train_granted[tgcnt]["plainText"][
                    i["start"] : i["end"]
                ]
            tgcnt = tgcnt + 1
            anno_sent_dict = {k: v for k, v in sorted(anno_sent_dict.items())}
            docid_anno_dict[k] = list(anno_sent_dict.values())
            docid_anno_start_pos[k] = list(anno_sent_dict.keys())
            anno_sent_dict = {}

        with open("docid_anno_dict.pkl", "wb") as f:
            pickle.dump(docid_anno_dict, f)
        with open("docid_anno_start_pos.pkl", "wb") as f:
            pickle.dump(docid_anno_start_pos, f)

    print("Done")
    print("Creating Highlighted Docs...")
    if len(os.listdir(highlighted_doc_filepath_anno)) == 0:
        count = 0
        for k, v in docid_anno_dict.items():
            if count <= NUM_HIGHLIGHTED_DOCS:
                filepath = highlighted_doc_filepath_anno + str(k)
                generate_highlighted_doc(filepath, v, WD_COLOR_INDEX.GREEN)
                count = count + 1
    print("Finishes :)")
    error_dict = analyze_error(docid_anno_start_pos, docid_sent_start_pos)
    return docid_sent_dict, docid_sent_start_pos, docid_anno_dict, docid_anno_start_pos


def main():
    args = parser.parse_args()
    luima_unlabelled_sent_seg()
    # sent_dic, _, anno_dict, _ = standard_sent_seg()
    # luima_sbd_sent_seg()
    # spacy_sentences_per_doc = [len(v) for k, v in sent_dic.items()]
    # annotated_sentences_per_doc = [len(v) for k, v in anno_dict.items()]
    # plt.plot(spacy_sentences_per_doc, label="Predicted Sentences")
    # plt.plot(annotated_sentences_per_doc, label="Annotated Sentences")
    # plt.legend()
    # plt.show()


if __name__ == "__main__":
    main()
